"""
将 .docx 文档高保真转换为单个 Markdown 文件。
图片提取到资源目录，其余所有内容（标题、正文、表格、链接、图片引用）
全部收敛到一个 .md 文件中。

支持的链接/跳转类型：
  - 外部超链接 (w:hyperlink + r:id)  -> [text](url)
  - 内部书签跳转 (w:hyperlink + w:anchor) -> [text](#anchor)
  - 书签锚点 (w:bookmarkStart) -> <a id="anchor"></a>
  - 交叉引用域代码 (w:fldChar + w:instrText REF) -> [text](#anchor)
"""
import argparse
import os
import re
import zipfile
from pathlib import Path
from lxml import etree

from docx import Document
from docx.oxml.ns import qn
from docx.text.run import Run

NS = {
    'w':  'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r':  'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a':  'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'v':  'urn:schemas-microsoft-com:vml',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
}

# font-size (EMU) -> heading level, derived from format_chapter7.py
SIZE_HEADING_MAP = {
    203200: 1,   # 16pt  章标题
    190500: 2,   # 15pt  一级节标题 (7.x)
    177800: 3,   # 14pt  二级节标题 (一、二、...)
}
BODY_SIZE = 152400       # 12pt
CAPTION_SIZE = 133350    # 10.5pt

# Word 内置 Heading 样式名 -> Markdown 标题级别
STYLE_HEADING_MAP = {
    'Heading 1': 1, 'Heading 2': 2, 'Heading 3': 3,
    'Heading 4': 4, 'Heading 5': 5, 'Heading 6': 6,
    'heading 1': 1, 'heading 2': 2, 'heading 3': 3,
    'heading 4': 4, 'heading 5': 5, 'heading 6': 6,
    'Title': 1,
}

# _bookmarkStart 中需要忽略的系统书签
_SYSTEM_BOOKMARKS = {'_GoBack', '_Toc', '_Ref', '_Hlt'}


# ── image extraction ────────────────────────────────────

def extract_images(docx_path: str, assets_dir: str) -> dict:
    """从 docx 的 zip 中提取所有 media/ 图片，返回 {原始路径: 导出路径}。"""
    os.makedirs(assets_dir, exist_ok=True)
    mapping = {}
    with zipfile.ZipFile(docx_path, 'r') as zf:
        for name in zf.namelist():
            if name.startswith('word/media/'):
                fname = os.path.basename(name)
                out_path = os.path.join(assets_dir, fname)
                with zf.open(name) as src, open(out_path, 'wb') as dst:
                    dst.write(src.read())
                mapping[name] = out_path
    return mapping


def build_rel_image_map(doc: Document, image_files: dict, assets_dir: str, md_dir: str) -> dict:
    """建立 rId -> markdown 相对路径 的映射。"""
    rel_map = {}
    for rel_id, rel in doc.part.rels.items():
        if 'image' in str(rel.reltype):
            target = f'word/{rel.target_ref}'
            if target in image_files:
                abs_path = image_files[target]
                rel_path = os.path.relpath(abs_path, md_dir).replace('\\', '/')
                rel_map[rel_id] = rel_path
    return rel_map


# ── bookmark collection ─────────────────────────────────

def _collect_bookmarks(body) -> dict:
    """扫描 body 中所有 w:bookmarkStart，返回 {bookmark_name: bookmark_id}。
    过滤掉 Word 系统内部书签（_GoBack 等）。"""
    bookmarks = {}
    for bm in body.findall('.//w:bookmarkStart', NS):
        name = bm.get(qn('w:name'), '')
        bm_id = bm.get(qn('w:id'), '')
        if name and not any(name.startswith(prefix) for prefix in _SYSTEM_BOOKMARKS):
            bookmarks[name] = bm_id
    return bookmarks


def _get_bookmarks_in_para(para_elem) -> list:
    """获取段落中的 bookmarkStart 名称列表（直接子元素）。"""
    names = []
    for child in para_elem:
        if etree.QName(child.tag).localname == 'bookmarkStart':
            name = child.get(qn('w:name'), '')
            if name and not any(name.startswith(p) for p in _SYSTEM_BOOKMARKS):
                names.append(name)
    return names


# ── paragraph-level helpers ─────────────────────────────

def _get_font_size(run_elem):
    """从 run XML 获取字号 (EMU)，优先 w:sz。"""
    rPr = run_elem.find('w:rPr', NS)
    if rPr is not None:
        sz = rPr.find('w:sz', NS)
        if sz is not None:
            half_pt = int(sz.get(qn('w:val')))
            return half_pt * 6350  # half-point -> EMU
    return None


def _detect_heading_level(para) -> int:
    """检测标题级别。优先 Word Heading 样式，其次字号映射。返回 0 表示非标题。"""
    if not para.text.strip():
        return 0
    # 优先：Word 内置 Heading 样式
    style_name = para.style.name if para.style else ''
    if style_name in STYLE_HEADING_MAP:
        return STYLE_HEADING_MAP[style_name]
    # 其次：基于字号判断（兼容 format_chapter7.py 生成的文档）
    for run in para.runs:
        if not run.text.strip():
            continue
        size = run.font.size
        if size is None:
            size = _get_font_size(run._element)
        if size is not None and size in SIZE_HEADING_MAP:
            return SIZE_HEADING_MAP[size]
        break
    return 0


def _is_caption(para) -> bool:
    """判断是否为图题/表题（10.5pt 居中）。"""
    text = para.text.strip()
    if not text:
        return False
    if re.match(r'^[图表]\d+-\d+', text):
        return True
    for run in para.runs:
        size = run.font.size
        if size is None:
            size = _get_font_size(run._element)
        if size == CAPTION_SIZE:
            return True
        break
    return False


def _clean_alt_text(raw: str) -> str:
    """清理图片 alt text：去除文件路径、重复文件名等噪音。"""
    if not raw:
        return ''
    if re.match(r'^[A-Za-z]:[/\\]', raw) or raw.startswith('/'):
        basename = os.path.splitext(os.path.basename(raw.rstrip()))[0]
        m = re.search(r'(图\d+-\d+\s*.+)', basename)
        if m:
            return m.group(1).strip()
        return basename
    return raw


def _extract_images_from_para(para_elem, rel_map: dict) -> list:
    """从段落 XML 中提取所有图片引用，返回 [(rel_path, alt_text), ...]。"""
    images = []
    for drawing in para_elem.findall('.//w:drawing', NS):
        blip = drawing.find(f'.//{{{NS["a"]}}}blip')
        if blip is not None:
            embed = blip.get(qn('r:embed'))
            if embed and embed in rel_map:
                doc_pr = drawing.find(f'.//{{{NS["wp"]}}}docPr')
                alt = ''
                if doc_pr is not None:
                    raw = doc_pr.get('descr', '') or doc_pr.get('name', '')
                    alt = _clean_alt_text(raw)
                images.append((rel_map[embed], alt))
    for imagedata in para_elem.findall(f'.//{{{NS["v"]}}}imagedata'):
        rid = imagedata.get(qn('r:id'))
        if rid and rid in rel_map:
            images.append((rel_map[rid], ''))
    return images


def _build_inline_text(para, rel_map: dict, doc_rels: dict, bookmarks: dict) -> str:
    """将段落的 runs / hyperlinks / field codes 组装为 Markdown 行内文本。
    处理：加粗、斜体、外部链接、内部书签跳转、交叉引用域代码。"""
    parts = []
    in_field = False       # 是否在域代码区间内
    field_instr = ''       # 域指令文本
    field_display = ''     # 域显示文本

    for child in para._element:
        tag = etree.QName(child.tag).localname

        # ── 超链接 (外部 + 内部书签跳转) ──
        if tag == 'hyperlink':
            rid = child.get(qn('r:id'))
            anchor = child.get(qn('w:anchor'))
            link_text = ''.join(
                t.text or '' for t in child.findall('.//w:t', NS)
            )
            if rid and rid in doc_rels:
                parts.append(f'[{link_text}]({doc_rels[rid]})')
            elif anchor:
                parts.append(f'[{link_text}](#{anchor})')
            else:
                parts.append(link_text)
            continue

        if tag != 'r':
            continue

        # ── 域代码处理 (交叉引用 REF) ──
        fld_char = child.find('w:fldChar', NS)
        if fld_char is not None:
            fld_type = fld_char.get(qn('w:fldCharType'))
            if fld_type == 'begin':
                in_field = True
                field_instr = ''
                field_display = ''
                continue
            elif fld_type == 'separate':
                continue
            elif fld_type == 'end':
                in_field = False
                ref_match = re.search(r'REF\s+(\S+)', field_instr)
                if ref_match:
                    ref_target = ref_match.group(1)
                    if ref_target in bookmarks:
                        parts.append(f'[{field_display}](#{ref_target})')
                    else:
                        parts.append(field_display)
                else:
                    parts.append(field_display)
                continue

        instr_text = child.find('w:instrText', NS)
        if instr_text is not None and in_field:
            field_instr += instr_text.text or ''
            continue

        # 域显示文本（separate 和 end 之间的 run）
        if in_field:
            run_texts = child.findall('w:t', NS)
            for rt in run_texts:
                field_display += rt.text or ''
            continue

        # ── 普通 run ──
        run = Run(child, para)
        text = run.text
        if not text:
            continue
        bold = run.bold
        italic = run.italic
        if bold and italic:
            parts.append(f'***{text}***')
        elif bold:
            parts.append(f'**{text}**')
        elif italic:
            parts.append(f'*{text}*')
        else:
            parts.append(text)

    return ''.join(parts)


# ── table conversion ────────────────────────────────────

def _cell_text(cell, doc_rels: dict = None) -> str:
    """提取单元格文本，支持超链接。多段落用 <br> 连接。"""
    texts = []
    for p in cell.paragraphs:
        if doc_rels:
            # 检查段落中是否有超链接
            has_links = len(p._element.findall('w:hyperlink', NS)) > 0
            if has_links:
                parts = []
                for child in p._element:
                    tag = etree.QName(child.tag).localname
                    if tag == 'hyperlink':
                        rid = child.get(qn('r:id'))
                        link_text = ''.join(
                            t.text or '' for t in child.findall('.//w:t', NS)
                        )
                        if rid and rid in doc_rels:
                            parts.append(f'[{link_text}]({doc_rels[rid]})')
                        else:
                            anchor = child.get(qn('w:anchor'))
                            if anchor:
                                parts.append(f'[{link_text}](#{anchor})')
                            else:
                                parts.append(link_text)
                    elif tag == 'r':
                        run_texts = child.findall('w:t', NS)
                        for rt in run_texts:
                            parts.append(rt.text or '')
                texts.append(''.join(parts))
                continue
        texts.append(p.text.strip())
    return '<br>'.join(t for t in texts if t)


def _table_has_merged_cells(table) -> bool:
    """检测表格是否有合并单元格。"""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            grid_span = tc.find('.//w:gridSpan', NS)
            v_merge = tc.find('.//w:vMerge', NS)
            if grid_span is not None:
                val = grid_span.get(qn('w:val'))
                if val and int(val) > 1:
                    return True
            if v_merge is not None:
                return True
    return False


def _table_to_md(table, doc_rels: dict = None) -> str:
    """将 python-docx Table 转为 Markdown 表格或 HTML table。"""
    rows_data = []
    for row in table.rows:
        row_cells = [_cell_text(c, doc_rels) for c in row.cells]
        rows_data.append(row_cells)

    if not rows_data:
        return ''

    if _table_has_merged_cells(table):
        return _table_to_html(table, doc_rels)

    ncols = max(len(r) for r in rows_data)
    for r in rows_data:
        while len(r) < ncols:
            r.append('')

    # 去重相邻相同单元格（python-docx 对合并单元格会重复）
    for r in rows_data:
        seen = {}
        for ci in range(len(r)):
            if ci > 0 and r[ci] == r[ci - 1] and r[ci]:
                r[ci] = ''

    lines = []
    header = rows_data[0]
    lines.append('| ' + ' | '.join(header) + ' |')
    lines.append('| ' + ' | '.join(['---'] * ncols) + ' |')
    for row in rows_data[1:]:
        lines.append('| ' + ' | '.join(row) + ' |')
    return '\n'.join(lines)


def _table_to_html(table, doc_rels: dict = None) -> str:
    """复杂表格降级为 HTML。"""
    lines = ['<table>', '<thead>', '<tr>']
    first_row = True
    for row in table.rows:
        if not first_row:
            lines.append('<tr>')
        for cell in row.cells:
            tag = 'th' if first_row else 'td'
            text = _cell_text(cell, doc_rels)
            lines.append(f'  <{tag}>{text}</{tag}>')
        lines.append('</tr>')
        if first_row:
            lines.append('</thead>')
            lines.append('<tbody>')
            first_row = False
    lines.append('</tbody>')
    lines.append('</table>')
    return '\n'.join(lines)


# ── main conversion ────────────────────────────────────

def convert(docx_path: str, output_path: str, assets_dir: str):
    docx_path = os.path.abspath(docx_path)
    output_path = os.path.abspath(output_path)
    assets_dir = os.path.abspath(assets_dir)
    md_dir = os.path.dirname(output_path)
    os.makedirs(md_dir, exist_ok=True)

    image_files = extract_images(docx_path, assets_dir)
    doc = Document(docx_path)
    rel_map = build_rel_image_map(doc, image_files, assets_dir, md_dir)

    # 收集 hyperlink rels (外部链接)
    doc_rels = {}
    for rel_id, rel in doc.part.rels.items():
        if 'hyperlink' in str(rel.reltype).lower():
            doc_rels[rel_id] = rel.target_ref

    body = doc.element.body

    # 收集所有书签，用于交叉引用解析
    bookmarks = _collect_bookmarks(body)

    md_lines = []
    img_counter = 0

    para_index = 0
    table_index = 0

    for child in body:
        tag = etree.QName(child.tag).localname

        if tag == 'p':
            para = doc.paragraphs[para_index]
            para_index += 1
            text = para.text.strip()

            # 提取段落中的图片
            images = _extract_images_from_para(child, rel_map)
            if images:
                for img_path, alt_text in images:
                    img_counter += 1
                    alt = alt_text if alt_text else f'image_{img_counter:03d}'
                    md_lines.append(f'![{alt}]({img_path})')
                    md_lines.append('')
                # 含图片的段落中的文本通常来自绘图对象内的文本框，跳过
                continue

            if not text:
                continue

            # 判断标题
            heading_level = _detect_heading_level(para)
            if heading_level > 0:
                # 如果标题段落包含书签锚点，插入 HTML anchor
                bm_names = _get_bookmarks_in_para(child)
                anchor_html = ''
                for bm in bm_names:
                    anchor_html += f'<a id="{bm}"></a>'
                prefix = '#' * heading_level
                if anchor_html:
                    md_lines.append(f'{anchor_html}')
                    md_lines.append(f'{prefix} {text}')
                else:
                    md_lines.append(f'{prefix} {text}')
                md_lines.append('')
                continue

            # 判断图题/表题
            if _is_caption(para):
                md_lines.append(f'*{text}*')
                md_lines.append('')
                continue

            # 检查段落中是否有书签锚点（非标题段落）
            bm_names = _get_bookmarks_in_para(child)
            if bm_names:
                for bm in bm_names:
                    md_lines.append(f'<a id="{bm}"></a>')

            # 普通段落：组装行内格式（含链接、域代码）
            inline = _build_inline_text(para, rel_map, doc_rels, bookmarks)
            md_lines.append(inline)
            md_lines.append('')

        elif tag == 'tbl':
            table = doc.tables[table_index]
            table_index += 1
            md_lines.append(_table_to_md(table, doc_rels))
            md_lines.append('')

    # 清理多余空行（连续 3 个以上空行合并为 2 个）
    result = []
    blank_count = 0
    for line in md_lines:
        if line.strip() == '':
            blank_count += 1
            if blank_count <= 2:
                result.append('')
        else:
            blank_count = 0
            result.append(line)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(result).strip() + '\n')

    img_count = len(os.listdir(assets_dir)) if os.path.isdir(assets_dir) else 0
    print(f'Done -> {output_path}')
    if img_count:
        print(f'Images -> {assets_dir} ({img_count} files)')
    else:
        print(f'No images extracted.')


def main():
    parser = argparse.ArgumentParser(
        description='将 .docx 文档转换为单个 Markdown 文件'
    )
    parser.add_argument('--input', required=True, help='输入 .docx 文件路径')
    parser.add_argument('--output', default=None, help='输出 .md 文件路径')
    parser.add_argument('--assets-dir', default=None, help='图片资源输出目录')
    args = parser.parse_args()

    docx_path = args.input
    if not os.path.isfile(docx_path):
        print(f'Error: file not found: {docx_path}')
        return

    stem = Path(docx_path).stem
    out_dir = os.path.join(os.path.dirname(docx_path) or '.', 'output')
    output_path = args.output or os.path.join(out_dir, f'{stem}.md')
    assets_dir = args.assets_dir or os.path.join(os.path.dirname(output_path), 'images')

    convert(docx_path, output_path, assets_dir)


if __name__ == '__main__':
    main()
