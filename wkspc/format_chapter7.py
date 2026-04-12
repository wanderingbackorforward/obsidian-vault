"""
将 7.1_draft.md、7.2_draft.md、7.3_draft.md 合并并生成符合论文排版规范的 docx。
"""
import re, os, io
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont

WORK = r"D:\mine\mynotes\obsidian-vault\wkspc"
FILES = ["7.1_draft.md", "7.2_draft.md", "7.3_draft.md"]
OUT_DIR = f"{WORK}\\chapter7_output"        # 统一输出文件夹
OUTPUT = f"{OUT_DIR}\\chapter7_final.docx"
PLACEHOLDER_DIR = f"{OUT_DIR}\\placeholders"
FIGURE_LIST = f"{OUT_DIR}\\figure_list.txt"

# ── helpers ──────────────────────────────────────────────

# 收集所有图片信息，用于生成图名列表
figure_list = []

def create_placeholder_image(fig_id, fig_caption, width_cm=13, height_cm=8):
    """生成占位图：灰底白框，中间显示图号，保存为 PNG。
    返回图片文件路径。"""
    os.makedirs(PLACEHOLDER_DIR, exist_ok=True)
    dpi = 150
    w_px = int(width_cm / 2.54 * dpi)
    h_px = int(height_cm / 2.54 * dpi)

    img = Image.new('RGB', (w_px, h_px), color=(230, 230, 230))
    draw = ImageDraw.Draw(img)

    # 白色边框
    draw.rectangle([10, 10, w_px - 11, h_px - 11], outline=(180, 180, 180), width=2)

    # 中央文字（图号）
    label = fig_id  # e.g. "图7-1"
    try:
        font = ImageFont.truetype("msyh.ttc", 36)  # 微软雅黑
    except Exception:
        try:
            font = ImageFont.truetype("simsun.ttc", 36)
        except Exception:
            font = ImageFont.load_default()
    bbox = draw.textbbox((0, 0), label, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.text(((w_px - tw) / 2, (h_px - th) / 2 - 20), label, fill=(120, 120, 120), font=font)

    # 副标题（提示替换）
    hint = "[ 请替换为实际图片 ]"
    try:
        font_sm = ImageFont.truetype("msyh.ttc", 20)
    except Exception:
        try:
            font_sm = ImageFont.truetype("simsun.ttc", 20)
        except Exception:
            font_sm = ImageFont.load_default()
    bbox2 = draw.textbbox((0, 0), hint, font=font_sm)
    tw2 = bbox2[2] - bbox2[0]
    draw.text(((w_px - tw2) / 2, (h_px - th) / 2 + 30), hint, fill=(160, 160, 160), font=font_sm)

    filepath = os.path.join(PLACEHOLDER_DIR, f"{fig_id}.png")
    img.save(filepath, 'PNG')
    return filepath

def add_figure_placeholder(doc, fig_id, fig_caption, fig_filename):
    """插入占位图（居中，宽13cm）+ 图题。同时记录到 figure_list。"""
    # 记录图片信息
    figure_list.append({
        'id': fig_id,
        'caption': fig_caption,
        'original_file': fig_filename,
    })

    # 生成并插入占位图
    placeholder_path = create_placeholder_image(fig_id, fig_caption)
    para_img = doc.add_paragraph()
    para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para_img.add_run()
    run.add_picture(placeholder_path, width=Cm(13))
    # 图片段落：段前12磅，不设 exact 行距（让图片自然撑高）
    pPr = para_img._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(int(12 * 20)))
    spacing.set(qn('w:after'), str(int(0 * 20)))
    # lineRule = auto，让段落高度自适应图片
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:line'), str(240))  # 240 twips = 单倍行距基准

    # 插入图题（图下方）
    add_figure_caption(doc, fig_caption)

    return para_img
def set_run_font(run, cn_font="宋体", en_font="Times New Roman", size_pt=12, bold=False):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = en_font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)

def set_paragraph_spacing(para, before_pt=0, after_pt=0, line_spacing_pt=None):
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), str(int(before_pt * 20)))
    spacing.set(qn('w:after'), str(int(after_pt * 20)))
    if line_spacing_pt:
        spacing.set(qn('w:line'), str(int(line_spacing_pt * 20)))
        spacing.set(qn('w:lineRule'), 'exact')

def set_first_line_indent(para, chars=2):
    pPr = para._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(chars * 100))

def add_heading(doc, text, level):
    """level: 0=章标题, 1=一级(7.x), 2=二级(7.x.x or 一/二/三...)"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    if level == 0:  # 章标题：黑体 三号16pt 加粗 居中
        set_run_font(run, "黑体", "Times New Roman", 16, bold=True)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(para, before_pt=24, after_pt=18, line_spacing_pt=20)
    elif level == 1:  # 一级：黑体 小三15pt 不加粗 左对齐
        set_run_font(run, "黑体", "Times New Roman", 15, bold=False)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(para, before_pt=24, after_pt=6, line_spacing_pt=20)
    elif level == 2:  # 二级：黑体 四号14pt 不加粗 左对齐
        set_run_font(run, "黑体", "Times New Roman", 14, bold=False)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(para, before_pt=12, after_pt=6, line_spacing_pt=20)
    return para

def add_body(doc, text):
    """正文：宋体/TNR 小四12pt 首行缩进2字符 ~22pt行距，去除所有加粗标记"""
    para = doc.add_paragraph()
    # 去掉 markdown 加粗标记，正文不需要加粗
    clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    r = para.add_run(clean)
    set_run_font(r, "宋体", "Times New Roman", 12, bold=False)
    set_first_line_indent(para, 2)
    set_paragraph_spacing(para, before_pt=0, after_pt=0, line_spacing_pt=22)
    return para

def add_table_caption(doc, text):
    """表题：表上方 黑体 五号10.5pt 居中"""
    para = doc.add_paragraph()
    run = para.add_run(text.replace('**', ''))
    set_run_font(run, "黑体", "Times New Roman", 10.5, bold=False)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(para, before_pt=12, after_pt=6)
    return para

def add_figure_caption(doc, text):
    """图题：图下方 宋体 五号10.5pt 居中"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, "宋体", "Times New Roman", 10.5, bold=False)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(para, before_pt=6, after_pt=12)
    return para

def add_table_from_rows(doc, headers, rows):
    """插入格式化表格"""
    if not headers or not rows:
        return
    ncols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    # header
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h.strip())
        set_run_font(r, "黑体", "Times New Roman", 10.5, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # data
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            val = row[ci].strip() if ci < len(row) else ""
            r = p.add_run(val)
            set_run_font(r, "宋体", "Times New Roman", 10.5, bold=False)
    # 设置表格内段落行距
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_paragraph_spacing(p, line_spacing_pt=16)

# ── parse markdown ───────────────────────────────────────
def parse_md(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    return lines

def is_table_sep(line):
    return bool(re.match(r'^\s*[-:]+[-|\s:]+$', line.strip()))

def is_table_border(line):
    return bool(re.match(r'^\s*-{5,}', line.strip()))

def preprocess_lines(lines):
    """合并被 Pandoc 拆断的续行，但保留表格区域原样。
    表格区域定义：以 **表7-x 标题** 开始，到最后一个 ----- 边框结束。"""
    # 第一步：标记表格区域
    # 找到所有 ----- 边框行的位置
    border_indices = []
    for idx, l in enumerate(lines):
        if re.match(r'^\s*-{5,}', l.strip()):
            border_indices.append(idx)

    # 每个 pandoc 表格恰好有3条边框线（顶部、分隔、底部），按3个一组配对
    table_ranges = []
    i = 0
    while i + 2 < len(border_indices):
        start = border_indices[i]
        end = border_indices[i + 2]
        # 往上找表格标题
        caption_start = start
        for k in range(start - 1, max(start - 5, -1), -1):
            if re.match(r'^\*\*表\d+-\d+', lines[k].strip()):
                caption_start = k
                break
        table_ranges.append((caption_start, end))
        i += 3

    # 标记表格行
    table_flags = [False] * len(lines)
    for s, e in table_ranges:
        for idx in range(s, e + 1):
            table_flags[idx] = True

    # 第二步：非表格区域内，合并连续非空非特殊行
    result = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # 表格区域行直接保留
        if table_flags[i]:
            result.append(line)
            i += 1
            continue

        # 特殊行直接保留
        if (not stripped or
            stripped.startswith('#') or
            re.match(r'^!\[', stripped)):
            result.append(line)
            i += 1
            continue

        # 普通文本行：向前看合并续行
        merged = stripped
        i += 1
        while i < len(lines):
            if table_flags[i]:
                break
            next_stripped = lines[i].strip()
            if (not next_stripped or
                next_stripped.startswith('#') or
                re.match(r'^!\[', next_stripped)):
                break
            merged += next_stripped
            i += 1
        result.append(merged)

    return [l + '\n' for l in result]

def process_lines(doc, lines):
    lines = preprocess_lines(lines)
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # skip empty
        if not stripped:
            i += 1
            continue

        # handle image lines (![caption](path){#ref})  -- insert placeholder + caption
        if re.match(r'^!\[', stripped):
            # extract caption and filename
            m = re.match(r'!\[(.*?)\]\((.*?)\)', stripped)
            if m:
                cap = m.group(1)        # e.g. "图7-1 系统总体架构图"
                fig_file = m.group(2)   # e.g. "figures/fig7-1_system_architecture.png"
                # 提取图号（如 "图7-1"）
                fig_id_match = re.match(r'(图\d+-\d+)', cap)
                fig_id = fig_id_match.group(1) if fig_id_match else cap.split()[0]
                add_figure_placeholder(doc, fig_id, cap, fig_file)
            i += 1
            continue

        # headings
        if stripped.startswith('#'):
            m = re.match(r'^(#+)\s+(.*)', stripped)
            if m:
                hashes = len(m.group(1))
                text = m.group(2).strip()
                if hashes == 1:  # # 第7章
                    add_heading(doc, text, 0)
                elif hashes == 2:  # ## 7.x
                    add_heading(doc, text, 1)
                elif hashes >= 3:  # ### 二级
                    # 去掉 "一、" 这种前缀中的中文数字序号格式不变，直接用
                    add_heading(doc, text, 2)
            i += 1
            continue

        # table caption (bold line like **表7-1 xxx**)
        if re.match(r'^\*\*表\d+-\d+', stripped):
            add_table_caption(doc, stripped)
            i += 1
            continue

        # pandoc-style table (lines of dashes)
        if is_table_border(stripped):
            # collect table block
            i += 1  # skip top border
            if i >= len(lines):
                break
            # read header line(s)
            header_lines = []
            while i < len(lines) and not is_table_border(lines[i].strip()) and not is_table_sep(lines[i].strip()):
                header_lines.append(lines[i].rstrip('\n'))
                i += 1
            # skip separator
            if i < len(lines):
                # parse column positions from the separator/border
                sep_line = lines[i].rstrip('\n')
                # find column boundaries from spaces in separator
                col_ranges = []
                in_col = False
                start = 0
                for ci, ch in enumerate(sep_line):
                    if ch == '-' and not in_col:
                        in_col = True
                        start = ci
                    elif ch == ' ' and in_col:
                        in_col = False
                        col_ranges.append((start, ci))
                if in_col:
                    col_ranges.append((start, len(sep_line)))
                i += 1

            if not col_ranges:
                col_ranges = None

            def display_width(s):
                """计算字符串的显示宽度（中文字符宽度为2）"""
                return sum(2 if ord(c) > 127 else 1 for c in s)

            def extract_cols_by_display_width(text_line, ranges):
                """按显示宽度位置切割内容行"""
                if not ranges:
                    return text_line.split()
                cols = []
                for col_start, col_end in ranges:
                    # 遍历字符，按显示宽度累加找到对应位置
                    char_start = None
                    char_end = None
                    dw = 0
                    for ci, ch in enumerate(text_line):
                        if dw >= col_start and char_start is None:
                            char_start = ci
                        if dw >= col_end:
                            char_end = ci
                            break
                        dw += 2 if ord(ch) > 127 else 1
                    if char_start is None:
                        cols.append("")
                    else:
                        if char_end is None:
                            char_end = len(text_line)
                        cols.append(text_line[char_start:char_end].strip())
                return cols

            # parse headers
            headers = extract_cols_by_display_width(header_lines[0] if header_lines else "", col_ranges)

            # read data rows until bottom border
            data_rows = []
            current_row = None
            while i < len(lines):
                dl = lines[i].rstrip('\n')
                if is_table_border(dl.strip()):
                    if current_row:
                        data_rows.append(current_row)
                        current_row = None
                    i += 1
                    break
                if not dl.strip():
                    # 空行作为行分隔符，结束当前行
                    if current_row:
                        data_rows.append(current_row)
                        current_row = None
                    i += 1
                    continue
                cols = extract_cols_by_display_width(dl, col_ranges)
                # check if this is a continuation (first col empty)
                if cols and cols[0] == '' and current_row:
                    # merge into current row
                    for ci in range(len(cols)):
                        if ci < len(current_row) and cols[ci]:
                            current_row[ci] = (current_row[ci] + " " + cols[ci]).strip()
                else:
                    if current_row:
                        data_rows.append(current_row)
                    current_row = cols
                i += 1
            if current_row:
                data_rows.append(current_row)

            add_table_from_rows(doc, headers, data_rows)
            continue

        # regular body paragraph - may span multiple lines until empty line or heading
        para_text = stripped
        i += 1
        # don't merge next lines, each non-empty non-special line is its own paragraph in md
        # Actually in these drafts, paragraphs are single long lines
        # Clean markdown artifacts
        para_text = re.sub(r'\{#fig:\S+\}', '', para_text)
        para_text = para_text.replace('------', '——')
        add_body(doc, para_text)
        continue

    return doc

# ── main ─────────────────────────────────────────────────
def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()

    # 页面设置 A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    for fname in FILES:
        path = f"{WORK}\\{fname}"
        lines = parse_md(path)
        process_lines(doc, lines)

    doc.save(OUTPUT)
    print(f"Done → {OUTPUT}")

    # 生成图名列表，方便用户查找和替换占位图
    with open(FIGURE_LIST, 'w', encoding='utf-8') as f:
        f.write("=" * 70 + "\n")
        f.write("  第7章 图片列表 — 占位图替换指南\n")
        f.write("=" * 70 + "\n\n")
        f.write("在Word中双击占位图即可替换为实际图片。\n")
        f.write("占位图文件保存在: placeholders\\ 文件夹中。\n\n")
        f.write(f"{'序号':<6}{'图号':<10}{'图题':<35}{'原始文件名'}\n")
        f.write("-" * 70 + "\n")
        for idx, fig in enumerate(figure_list, 1):
            f.write(f"{idx:<6}{fig['id']:<10}{fig['caption']:<35}{fig['original_file']}\n")
        f.write("-" * 70 + "\n")
        f.write(f"\n共 {len(figure_list)} 张图片需要替换。\n")
    print(f"Figure list → {FIGURE_LIST}")

if __name__ == "__main__":
    main()
