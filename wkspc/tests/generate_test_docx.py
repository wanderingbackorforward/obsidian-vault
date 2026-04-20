"""
生成一个最小测试 docx，专门覆盖链接与跳转相关场景。
用于验证 docx_to_md.py 对以下元素的处理能力：
  1. 外部超链接 (w:hyperlink + r:id)
  2. 内部书签跳转 (w:hyperlink + w:anchor / w:bookmarkStart)
  3. 交叉引用 / 域代码 (w:fldSimple / w:instrText)
  4. 混排链接（同一段落中链接与普通文本交替）
  5. 基础 Word Heading 样式
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.join(os.path.dirname(__file__), 'fixtures')
OUTPUT = os.path.join(OUT_DIR, 'docx_links_and_bookmarks.docx')


def add_hyperlink(paragraph, text, url):
    """向段落添加外部超链接。"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    hyperlink.append(run)
    paragraph._element.append(hyperlink)


def add_bookmark_start(paragraph, bookmark_id, bookmark_name):
    """在段落中插入 bookmarkStart。"""
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bookmark_id))
    bm_start.set(qn('w:name'), bookmark_name)
    paragraph._element.insert(0, bm_start)


def add_bookmark_end(paragraph, bookmark_id):
    """在段落中插入 bookmarkEnd。"""
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bookmark_id))
    paragraph._element.append(bm_end)


def add_internal_hyperlink(paragraph, text, anchor_name):
    """向段落添加内部书签跳转链接 (w:hyperlink w:anchor=...)。"""
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor_name)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    hyperlink.append(run)
    paragraph._element.append(hyperlink)


def add_cross_reference_field(paragraph, field_code, display_text):
    """插入域代码形式的交叉引用 (w:fldChar + w:instrText)。"""
    # fldChar begin
    fld_begin = OxmlElement('w:r')
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    fld_begin.append(fld_char_begin)
    paragraph._element.append(fld_begin)

    # instrText
    instr_run = OxmlElement('w:r')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = field_code
    instr_run.append(instr_text)
    paragraph._element.append(instr_run)

    # fldChar separate
    fld_sep = OxmlElement('w:r')
    fld_char_sep = OxmlElement('w:fldChar')
    fld_char_sep.set(qn('w:fldCharType'), 'separate')
    fld_sep.append(fld_char_sep)
    paragraph._element.append(fld_sep)

    # display text
    display_run = OxmlElement('w:r')
    dt = OxmlElement('w:t')
    dt.text = display_text
    dt.set(qn('xml:space'), 'preserve')
    display_run.append(dt)
    paragraph._element.append(display_run)

    # fldChar end
    fld_end = OxmlElement('w:r')
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    fld_end.append(fld_char_end)
    paragraph._element.append(fld_end)


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()

    # ── 第1节：标题 + 书签锚点 ──
    h1 = doc.add_heading('第1节 测试文档概述', level=1)
    add_bookmark_start(h1, 1, 'section_1')
    add_bookmark_end(h1, 1)

    p = doc.add_paragraph('本文档用于测试 docx_to_md.py 对链接与跳转的处理能力。')

    # 外部超链接：独立段落
    p2 = doc.add_paragraph('请访问 ')
    add_hyperlink(p2, 'OpenAI 官网', 'https://openai.com/')
    p2.add_run(' 了解更多信息。')

    # 外部超链接：多个链接混排
    p3 = doc.add_paragraph('相关资源：')
    add_hyperlink(p3, 'GitHub', 'https://github.com/')
    p3.add_run('、')
    add_hyperlink(p3, 'Python 文档', 'https://docs.python.org/3/')
    p3.add_run('、')
    add_hyperlink(p3, 'Stack Overflow', 'https://stackoverflow.com/')
    p3.add_run('。')

    # 内部跳转：跳到第2节
    p4 = doc.add_paragraph('关于链接测试的详细内容，请参见')
    add_internal_hyperlink(p4, '第2节 外部超链接测试', 'section_2')
    p4.add_run('。')

    # 内部跳转：跳到第3节
    p5 = doc.add_paragraph('关于书签测试，请参见')
    add_internal_hyperlink(p5, '第3节', 'section_3')
    p5.add_run('。')

    # ── 第2节：外部超链接 ──
    h2 = doc.add_heading('第2节 外部超链接测试', level=2)
    add_bookmark_start(h2, 2, 'section_2')
    add_bookmark_end(h2, 2)

    doc.add_paragraph('本节测试外部超链接在转换后是否保留为 Markdown 链接格式。')

    p6 = doc.add_paragraph('这是一个')
    p6.add_run('加粗文本').bold = True
    p6.add_run('后面紧跟')
    add_hyperlink(p6, '一个链接', 'https://example.com/test')
    p6.add_run('的混排段落。')

    # ── 第3节：内部书签跳转 ──
    h3 = doc.add_heading('第3节 内部书签跳转测试', level=2)
    add_bookmark_start(h3, 3, 'section_3')
    add_bookmark_end(h3, 3)

    doc.add_paragraph('本节测试内部书签跳转（w:anchor）在转换后的处理。')

    p7 = doc.add_paragraph('点击可以跳回')
    add_internal_hyperlink(p7, '第1节 测试文档概述', 'section_1')
    p7.add_run('。')

    # 在段落中设置一个独立书签
    p8 = doc.add_paragraph('这里有一个重要标记。')
    add_bookmark_start(p8, 4, 'important_mark')
    add_bookmark_end(p8, 4)

    p9 = doc.add_paragraph('从这里跳转到')
    add_internal_hyperlink(p9, '重要标记位置', 'important_mark')
    p9.add_run('。')

    # ── 第4节：交叉引用 ──
    h4 = doc.add_heading('第4节 交叉引用与域代码测试', level=2)
    add_bookmark_start(h4, 5, 'section_4')
    add_bookmark_end(h4, 5)

    doc.add_paragraph('本节测试 Word 域代码（field code）形式的交叉引用。')

    p10 = doc.add_paragraph('如')
    add_cross_reference_field(p10, ' REF section_2 \\h ', '第2节 外部超链接测试')
    p10.add_run('所述，外部链接应保留完整 URL。')

    p11 = doc.add_paragraph('另见')
    add_cross_reference_field(p11, ' REF important_mark \\h ', '重要标记位置')
    p11.add_run('中的说明。')

    # ── 第5节：简单表格 + 表内链接 ──
    h5 = doc.add_heading('第5节 表格与链接混合', level=2)

    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '名称'
    table.rows[0].cells[1].text = '链接'
    table.rows[1].cells[0].text = 'Python'
    # 表格单元格内的链接
    cell_para = table.rows[1].cells[1].paragraphs[0]
    add_hyperlink(cell_para, 'python.org', 'https://www.python.org/')
    table.rows[2].cells[0].text = 'Rust'
    cell_para2 = table.rows[2].cells[1].paragraphs[0]
    add_hyperlink(cell_para2, 'rust-lang.org', 'https://www.rust-lang.org/')

    # ── 结尾 ──
    doc.add_paragraph('')
    p_end = doc.add_paragraph('返回')
    add_internal_hyperlink(p_end, '文档开头', 'section_1')
    p_end.add_run('。')

    doc.save(OUTPUT)
    print(f'Test fixture generated -> {OUTPUT}')


if __name__ == '__main__':
    main()
